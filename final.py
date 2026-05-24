import streamlit as st
import PyPDF2
import docx
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import plotly.graph_objects as go
from collections import Counter
import io
import time
from datetime import datetime

st.set_page_config(page_title="ResumeIQ", page_icon="📄", layout="wide", initial_sidebar_state="collapsed")

# ═══════════════════════════════════════════════════════
#  STYLES
# ═══════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&family=Syne:wght@700;800&display=swap');

html, body, [data-testid="stAppViewContainer"] {
    background:#0d0d16 !important; font-family:'Inter',sans-serif !important; color:#e2e8f0 !important;
}
[data-testid="stSidebar"] { background:#09090f !important; border-right:1px solid #1e1e30 !important; }
[data-testid="stSidebar"] label,[data-testid="stSidebar"] p,[data-testid="stSidebar"] span { color:#94a3b8 !important; }

/* LANDING */
.landing-wrap {
    min-height:80vh; display:flex; flex-direction:column;
    align-items:center; justify-content:center; text-align:center; padding:60px 20px;
}
.landing-chip {
    display:inline-block; background:rgba(139,92,246,.18); border:1px solid rgba(139,92,246,.4);
    color:#c4b5fd; font-size:.72rem; font-weight:700; letter-spacing:1.5px;
    padding:5px 18px; border-radius:20px; margin-bottom:24px;
}
.landing-title {
    font-family:'Syne',sans-serif; font-size:4rem; font-weight:800;
    background:linear-gradient(90deg,#a78bfa 0%,#34d399 100%);
    -webkit-background-clip:text; -webkit-text-fill-color:transparent; background-clip:text;
    margin:0 0 16px; line-height:1.1;
}
.landing-sub { color:#94a3b8; font-size:1.1rem; font-weight:300; max-width:520px; margin:0 auto 36px; line-height:1.7; }
.landing-features {
    display:flex; gap:24px; flex-wrap:wrap; justify-content:center; margin-top:48px;
}
.landing-feat {
    background:#11111e; border:1px solid #1e1e30; border-radius:14px;
    padding:20px 24px; width:160px; text-align:center;
}
.landing-feat-icon { font-size:1.8rem; margin-bottom:8px; }
.landing-feat-label { font-size:.78rem; color:#94a3b8; }

/* HERO */
.hero {
    background:linear-gradient(135deg,#0f1117 0%,#151c2e 50%,#0f1117 100%);
    border:1px solid #2a1a4a; border-radius:22px;
    padding:40px 40px 34px; text-align:center; margin-bottom:26px;
    position:relative; overflow:hidden;
}
.hero::before { content:'';position:absolute;top:-80px;right:-80px;width:300px;height:300px;background:radial-gradient(circle,rgba(139,92,246,.22) 0%,transparent 68%);border-radius:50%; }
.hero::after  { content:'';position:absolute;bottom:-60px;left:15%;width:240px;height:240px;background:radial-gradient(circle,rgba(16,185,129,.15) 0%,transparent 68%);border-radius:50%; }
.hero-chip { display:inline-block;background:rgba(139,92,246,.18);border:1px solid rgba(139,92,246,.4);color:#c4b5fd;font-size:.7rem;font-weight:700;letter-spacing:1.5px;padding:4px 16px;border-radius:20px;margin-bottom:14px; }
.hero h1 { font-family:'Syne',sans-serif;font-size:2.6rem;font-weight:800;background:linear-gradient(90deg,#a78bfa,#34d399);-webkit-background-clip:text;-webkit-text-fill-color:transparent;background-clip:text;margin:0 0 10px; }
.hero-sub { color:#94a3b8;font-size:.92rem;font-weight:300;margin:0 0 24px; }
.hero-stats { display:flex;justify-content:center;gap:40px;flex-wrap:wrap;padding-top:22px;border-top:1px solid rgba(255,255,255,.07); }
.hero-stat-n { font-family:'Syne',sans-serif;font-size:1.5rem;font-weight:800;background:linear-gradient(90deg,#a78bfa,#34d399);-webkit-background-clip:text;-webkit-text-fill-color:transparent;background-clip:text; }
.hero-stat-l { color:#475569;font-size:.72rem;margin-top:2px; }

/* CARDS */
.card { background:#11111e;border:1px solid #1e1e30;border-radius:16px;padding:22px 24px;margin-bottom:18px; }
.sec-title { font-family:'Syne',sans-serif;font-size:1rem;font-weight:700;color:#c4b5fd;margin-bottom:14px; }

/* SKILL GRID */
.skill-cat-card { background:#0d0d1a;border:1px solid #1e1e30;border-radius:12px;padding:14px 16px;margin-bottom:10px; }
.skill-cat-title { font-size:.75rem;font-weight:700;color:#7c3aed;letter-spacing:.5px;text-transform:uppercase;margin-bottom:8px; }
.skill-mini-prog-wrap { background:#1e2a3e;border-radius:6px;height:5px;overflow:hidden;margin:6px 0 10px; }
.skill-mini-prog { height:100%;border-radius:6px; }

/* PILLS */
.pill { display:inline-block;padding:4px 12px;border-radius:20px;font-size:.77rem;margin:3px;font-weight:500; }
.have    { background:#052e16;color:#6ee7b7;border:1px solid #065f46; }
.missing { background:#1f0a0a;color:#fca5a5;border:1px solid #7f1d1d; }
.neutral { background:#0f172a;color:#93c5fd;border:1px solid #1e3a5f; }

/* BADGES */
.badge { display:inline-block;padding:3px 9px;border-radius:10px;font-size:.68rem;font-weight:800;margin:2px 4px 2px 0;letter-spacing:.4px; }
.badge-high   { background:#2d0a0a;color:#f87171;border:1px solid #7f1d1d; }
.badge-medium { background:#1c1200;color:#fbbf24;border:1px solid #713f12; }
.badge-low    { background:#0c1a2e;color:#93c5fd;border:1px solid #1e3a5f; }

/* TIP BOXES */
.tip-box       { background:#0f0f1e;border-left:3px solid #7c3aed;padding:13px 16px;border-radius:8px;margin:8px 0;font-size:.88rem;color:#cbd5e1;line-height:1.65; }
.tip-box-green { background:#051a05;border-left:3px solid #059669;padding:13px 16px;border-radius:8px;margin:8px 0;font-size:.88rem;color:#d1fae5;line-height:1.65; }
.tip-box-red   { background:#1a0505;border-left:3px solid #dc2626;padding:13px 16px;border-radius:8px;margin:8px 0;font-size:.88rem;color:#fca5a5;line-height:1.65; }

/* ROAST */
.roast-box { background:linear-gradient(135deg,#1a0505,#1a0a00);border:1px solid #7f1d1d;border-radius:14px;padding:20px 24px;margin:10px 0;font-size:.92rem;color:#fca5a5;line-height:1.8; }

/* PROGRESS */
.prog-wrap { background:#1e1e2e;border-radius:8px;height:8px;overflow:hidden;margin:8px 0; }
.prog-fill { height:100%;border-radius:8px; }

/* CHATBOT */
.chat-wrap { background:#09090f;border:1px solid #1e1e30;border-radius:16px;padding:16px;max-height:440px;overflow-y:auto;margin:16px 0; }
.msg-row-user { display:flex;justify-content:flex-end;margin:10px 0;align-items:flex-end;gap:8px; }
.msg-row-bot  { display:flex;justify-content:flex-start;margin:10px 0;align-items:flex-end;gap:8px; }
.bubble-user { background:linear-gradient(135deg,#4c1d95,#1e3a5f);color:#e0e7ff;padding:11px 15px;border-radius:18px 18px 4px 18px;max-width:75%;font-size:.88rem;line-height:1.55;box-shadow:0 2px 12px rgba(124,58,237,.2); }
.bubble-bot  { background:#111827;border:1px solid #1e2d20;color:#d1fae5;padding:11px 15px;border-radius:18px 18px 18px 4px;max-width:75%;font-size:.88rem;line-height:1.55; }
.avatar { width:30px;height:30px;border-radius:50%;display:flex;align-items:center;justify-content:center;font-size:14px;flex-shrink:0; }
.avatar-user { background:#4c1d95; }
.avatar-bot  { background:#052e16; }
.typing-dot { display:inline-block;width:7px;height:7px;background:#34d399;border-radius:50%;margin:0 2px;animation:blink 1.2s infinite; }
.typing-dot:nth-child(2){animation-delay:.2s}.typing-dot:nth-child(3){animation-delay:.4s}
@keyframes blink{0%,80%,100%{opacity:.2}40%{opacity:1}}

/* WORD CLOUD */
.wcloud-word { display:inline-block;margin:4px;padding:5px 11px;border-radius:8px;font-weight:600;cursor:default;transition:transform .2s; }
.wcloud-word:hover { transform:scale(1.1); }

/* TEMPLATE CARDS */
.tmpl-card { background:#0d0d1a;border:2px solid #1e1e30;border-radius:14px;padding:20px;text-align:center;transition:border-color .2s; }
.tmpl-card:hover { border-color:#7c3aed; }
.tmpl-selected { border-color:#7c3aed !important;background:#170d2e !important; }

/* INTERVIEW Q */
.iq-card { background:#0d0d1a;border:1px solid #1e2d3a;border-radius:12px;padding:14px 18px;margin:8px 0;font-size:.88rem;color:#bae6fd;line-height:1.6; }
.iq-num  { font-family:'Syne',sans-serif;font-weight:800;color:#7c3aed;margin-right:8px; }

/* INPUTS & BUTTONS */
.stTextInput input,.stTextArea textarea,.stSelectbox select { background:#11111e !important;border:1px solid #2a2a3e !important;border-radius:10px !important;color:#e2e8f0 !important; }
.stTextInput label,.stTextArea label,.stSelectbox label { color:#c4b5fd !important;font-weight:600 !important;font-size:0.88rem !important; }
.stButton button, .stDownloadButton button { background:linear-gradient(135deg,#7c3aed,#059669) !important;color:white !important;border:none !important;border-radius:10px !important;font-weight:600 !important; }
[data-testid="stMetricValue"] { color:#a78bfa !important; font-size:1.8rem !important; }
[data-testid="stMetricLabel"] { color:#94a3b8 !important; }
hr { border-color:#1e1e30 !important; }
::-webkit-scrollbar{width:4px}::-webkit-scrollbar-track{background:#09090f}::-webkit-scrollbar-thumb{background:#2d2d4e;border-radius:4px}
</style>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════
#  DATA
# ═══════════════════════════════════════════════════════
SKILLS = {
    "Languages":  ["python","java","javascript","c++","typescript","go","ruby","swift"],
    "Web":        ["react","html","css","node.js","angular","vue","django","flask"],
    "Data & ML":  ["pandas","numpy","tensorflow","pytorch","scikit-learn","sql","keras","nlp"],
    "Cloud":      ["aws","azure","gcp","docker","kubernetes","git","linux","jenkins"],
}

ROLES = {
    "💻 Software Engineer":  ["python","java","c++","git"],
    "🎨 Frontend Developer": ["javascript","react","html","css"],
    "📊 Data Analyst":       ["sql","pandas","numpy","python"],
    "🧠 Data Scientist":     ["python","pandas","scikit-learn","tensorflow"],
    "☁️ DevOps Engineer":    ["docker","kubernetes","aws","git"],
}

COURSES = {
    "python":     ("Python for Everybody","Coursera","https://coursera.org"),
    "react":      ("React Complete Guide","Udemy","https://udemy.com"),
    "sql":        ("SQL for Data Science","Coursera","https://coursera.org"),
    "docker":     ("Docker Mastery","Udemy","https://udemy.com"),
    "tensorflow": ("Deep Learning Specialization","Coursera","https://coursera.org"),
    "javascript": ("JavaScript Algorithms","freeCodeCamp","https://freecodecamp.org"),
    "kubernetes": ("Kubernetes for Beginners","Linux Foundation","https://linuxfoundation.org"),
    "aws":        ("AWS Cloud Practitioner","AWS Training","https://aws.amazon.com/training"),
}

INTERVIEW_QUESTIONS = {
    "python":      ["What are Python decorators and when would you use them?","Explain the difference between list, tuple, and set in Python.","How does Python handle memory management?"],
    "java":        ["What is the difference between abstract class and interface in Java?","Explain JVM, JRE, and JDK.","What are Java generics?"],
    "javascript":  ["What is event bubbling and event capturing?","Explain the difference between == and === in JavaScript.","What is the event loop in JavaScript?"],
    "react":       ["What is the virtual DOM and how does React use it?","Explain useState vs useEffect hooks.","What are React keys and why are they important?"],
    "sql":         ["What is the difference between INNER JOIN and LEFT JOIN?","Explain database normalization.","What are indexes and how do they improve performance?"],
    "docker":      ["What is the difference between a Docker image and container?","Explain Docker Compose.","How do you persist data in Docker?"],
    "aws":         ["What is the difference between EC2 and Lambda?","Explain S3 and its storage classes.","What is IAM and how does it work?"],
    "machine learning": ["What is overfitting and how do you prevent it?","Explain bias-variance tradeoff.","What is cross-validation?"],
    "tensorflow":  ["What is a tensor?","Explain the difference between eager and graph execution.","What are callbacks in TensorFlow?"],
    "pandas":      ["What is the difference between loc and iloc?","How do you handle missing values in pandas?","What is groupby and how do you use it?"],
    "git":         ["What is the difference between merge and rebase?","Explain git stash.","How do you resolve a merge conflict?"],
    "kubernetes":  ["What is a pod in Kubernetes?","Explain the difference between Deployment and StatefulSet.","What is a Kubernetes Service?"],
}

ROAST_LINES = {
    "no_email":    "📧 No email? Were you planning to telepathically send your application?",
    "no_phone":    "📵 No phone number. Bold strategy — let's see if recruiters learn morse code.",
    "no_projects": "🛠 No projects section. So you learned to code but never actually... coded anything?",
    "no_percent":  "📊 Zero measurable results. 'Worked on stuff' is not an achievement. What did you actually accomplish?",
    "too_short":   "📄 This resume is shorter than a tweet. Recruiters spend 6 seconds on a resume — yours would take 2.",
    "no_skills":   "🔧 Skills section is either missing or a desert. 'Microsoft Office' doesn't count in 2025.",
    "no_github":   "🐙 No GitHub link. How is a recruiter supposed to know you wrote more than hello world?",
    "low_ats":     "🤖 Your ATS score is so low the algorithm probably marked it as spam.",
    "buzzwords":   "💼 Pro tip: 'Passionate', 'hardworking', 'team player' — every single resume says this. It means nothing.",
    "generic":     "😐 This resume is so generic it could belong to literally anyone. Make it yours.",
}

SAMPLE_DATA = {
    "name":"Rahul Sharma","email":"rahul.sharma@gmail.com","phone":"+91 98765 43210",
    "linkedin":"linkedin.com/in/rahulsharma","github":"github.com/rahulsharma",
    "summary":"Final year B.Tech CSE student with experience in Python, Machine Learning, and web development. Built 3+ end-to-end projects. Seeking a Data Scientist role at a product-driven company.",
    "skills":"Python, SQL, TensorFlow, scikit-learn, Pandas, NumPy, React, Docker, Git, Linux, REST APIs",
    "experience":"Machine Learning Intern — DataWorks AI (June–Aug 2024)\n• Developed text classification model using BERT, achieving 94% accuracy on 10K samples\n• Automated data pipeline reducing preprocessing time by 40%\n• Collaborated with 5 engineers using Git and Agile sprints",
    "education":"B.Tech Computer Science — VIT University (2021–2025) | CGPA: 8.6/10\nClass 12 — Delhi Public School (2021) | 94.2%",
    "projects":"Resume Analyzer (Python, Streamlit, NLP)\n• Built ATS resume checker used by 200+ students\n• Deployed on Streamlit Cloud with real-time PDF parsing\n\nStock Price Predictor (Python, LSTM, TensorFlow)\n• Predicted stock prices with 87% directional accuracy\n• Visualised with Plotly interactive charts",
    "certifications":"Deep Learning Specialization — Coursera (2024)\nAWS Cloud Practitioner — AWS Training (2023)",
}

# ═══════════════════════════════════════════════════════
#  UTILS
# ═══════════════════════════════════════════════════════

def extract_text(file) -> str:
    try:
        if file.type == "application/pdf":
            reader = PyPDF2.PdfReader(file)
            text = " ".join(page.extract_text() or "" for page in reader.pages)
            if len(text.strip()) < 50:
                return "[ERROR: Scanned PDF — please upload a text-based PDF.]"
            return text
        else:
            d = docx.Document(file)
            return " ".join(p.text for p in d.paragraphs)
    except Exception as e:
        return f"[ERROR: {e}]"


def get_skills(text: str) -> dict:
    t = text.lower()
    return {cat: {"found":[s for s in sk if s in t], "missing":[s for s in sk if s not in t]}
            for cat, sk in SKILLS.items()}


def run_ats(text: str):
    checks = [
        (bool(re.search(r'[\w.\-]+@[\w.\-]+\.\w+', text)), "Add your email address",           15),
        (bool(re.search(r'(\+?\d[\d\s\-().]{7,}\d)', text)),"Add a phone number",              10),
        ("experience" in text.lower(),                       "Add an 'Experience' section",    20),
        ("education"  in text.lower(),                       "Add an 'Education' section",     15),
        ("skills"     in text.lower(),                       "Add a 'Skills' section",         15),
        ("project"    in text.lower(),                       "Add a 'Projects' section",       10),
        ("%" in text,                                        "Add measurable results using %", 10),
        (len(text.split()) >= 200,                           "Resume too short — aim for 300+ words", 5),
    ]
    issues = [(msg, pen) for ok, msg, pen in checks if not ok]
    return max(0, 100 - sum(p for _, p in issues)), issues


def compute_score(text: str) -> int:
    skill_data = get_skills(text)
    ats_sc, _  = run_ats(text)
    found  = sum(len(d["found"]) for d in skill_data.values())
    total  = sum(len(d["found"])+len(d["missing"]) for d in skill_data.values())
    skill_pct = int(found/max(total,1)*100)
    return int(ats_sc*0.5 + skill_pct*0.5)


def prog_color(s):
    if s >= 70: return "linear-gradient(90deg,#059669,#34d399)"
    if s >= 40: return "linear-gradient(90deg,#d97706,#fbbf24)"
    return "linear-gradient(90deg,#dc2626,#f87171)"


def smart_feedback(text, skill_data, _ats):
    tips, tl = [], text.lower()
    if "%" not in text:          tips.append("📈 Add measurable achievements — e.g. 'Improved speed by 35%'. Numbers make you stand out.")
    if not re.search(r'[\w.\-]+@[\w.\-]+\.\w+', text): tips.append("📧 Add your email address — recruiters can't contact you without it.")
    if "linkedin" not in tl and "github" not in tl:    tips.append("🔗 Add LinkedIn and GitHub URLs — recruiters always check GitHub for fresher candidates.")
    if "project" not in tl:      tips.append("🛠 Add a Projects section — for freshers this is the most important section.")
    if not any(v in tl for v in ["developed","built","led","designed","implemented","deployed"]): tips.append("💪 Start bullets with action verbs — 'Developed', 'Built', 'Designed'. Avoid 'was responsible for'.")
    if len(text.split()) < 300:  tips.append(f"📝 Resume is only ~{len(text.split())} words. Expand project descriptions — aim for 400–600 words.")
    if sum(len(d['found']) for d in skill_data.values()) < 5: tips.append("🔧 Very few skills detected. Add a dedicated Skills section.")
    if not tips:
        tips.append("✅ Resume looks solid! Tailor it per job by mirroring keywords from each JD.")
        tips.append("🎯 Add a 2-line professional summary at the top — a quick pitch of who you are.")
    return tips[:5]


def generate_roast(text: str, skill_data: dict, ats_sc: int) -> list:
    """Brutally honest (but fun) feedback."""
    tl    = text.lower()
    roast = []
    if not re.search(r'[\w.\-]+@[\w.\-]+\.\w+', text): roast.append(ROAST_LINES["no_email"])
    if not re.search(r'(\+?\d[\d\s\-().]{7,}\d)', text): roast.append(ROAST_LINES["no_phone"])
    if "project" not in tl:  roast.append(ROAST_LINES["no_projects"])
    if "%" not in text:       roast.append(ROAST_LINES["no_percent"])
    if len(text.split()) < 250: roast.append(ROAST_LINES["too_short"])
    if "github" not in tl:   roast.append(ROAST_LINES["no_github"])
    if ats_sc < 50:           roast.append(ROAST_LINES["low_ats"])
    for word in ["passionate","hardworking","team player","go-getter","self-motivated"]:
        if word in tl: roast.append(ROAST_LINES["buzzwords"]); break
    if sum(len(d["found"]) for d in skill_data.values()) < 4: roast.append(ROAST_LINES["no_skills"])
    if not roast:
        roast.append("😤 Honestly? This resume is pretty decent. I'm struggling to roast it. Add more projects and % metrics and you'll be solid.")
    return roast[:5]


def get_interview_questions(skill_data: dict) -> list:
    """Return interview questions based on detected skills."""
    found  = [s for d in skill_data.values() for s in d["found"]]
    result = []
    for skill in found:
        if skill in INTERVIEW_QUESTIONS:
            for q in INTERVIEW_QUESTIONS[skill][:2]:
                result.append((skill, q))
        if len(result) >= 10: break
    if not result:
        result = [("general","Tell me about yourself."),("general","What are your strengths and weaknesses?"),
                  ("general","Where do you see yourself in 5 years?"),("general","Why should we hire you?")]
    return result


def smart_jd_analysis(matched, missing, score, jd):
    tips, jl = [], jd.lower()
    if score < 40:   tips.append(f"⚠️ Low match ({score}%) — tailor your resume specifically for this role before applying.")
    elif score < 70: tips.append(f"📊 Moderate match ({score}%) — decent base but add the missing keywords to improve chances.")
    else:            tips.append(f"✅ Strong match ({score}%) — your resume aligns well with this JD!")
    if missing: tips.append(f"🔑 Top keywords to add: **{', '.join(missing[:6])}** — add these to your Skills or Experience section.")
    if "leadership" in jl or "lead" in jl: tips.append("👥 JD mentions leadership — add examples where you led a team.")
    if "communication" in jl: tips.append("🗣 JD values communication — mention presentations or cross-team collaboration.")
    if "agile" in jl or "scrum" in jl: tips.append("🔄 JD mentions Agile/Scrum — if you've worked in sprints, add that.")
    likelihood = "High 🟢" if score >= 70 else ("Medium 🟡" if score >= 45 else "Low 🔴")
    tips.append(f"🎯 Interview likelihood: **{likelihood}** based on keyword match.")
    return tips


# ═══════════════════════════════════════════════════════
#  CHATBOT
# ═══════════════════════════════════════════════════════

def detect_intent(msg: str) -> str:
    m = msg.lower()
    intents = {
        "greeting":      ["hi","hello","hey","hii","sup","good morning","what's up"],
        "strengths":     ["strength","strong","good at","best at","what can i","i know","i have","my skills"],
        "weaknesses":    ["weak","missing","lack","gap","improve","not have","don't have","what should i learn"],
        "role_fit":      ["role","job","suited","which job","what job","career","fit for","good for","apply for","should i apply","recommend","google","amazon","microsoft","startup","faang"],
        "ats":           ["ats","applicant tracking","ats score","pass ats","automated","filter","shortlist"],
        "projects":      ["project","portfolio","github","showcase","build","side project"],
        "interview":     ["interview","prepare","preparation","interview tips","what to study","crack interview","get hired"],
        "salary":        ["salary","pay","ctc","package","lpa","how much","compensation","stipend","money"],
        "cover_letter":  ["cover","letter","cover letter"],
        "linkedin":      ["linkedin","profile","online presence","github link"],
        "fresher":       ["fresher","fresh graduate","no experience","first job","entry level","new grad","campus","placement"],
        "skills_list":   ["what skills","list skills","show skills","detect","found skills","which skills"],
        "action_verbs":  ["bullet","action verb","phrasing","how to write","write experience"],
        "summary":       ["summary","objective","profile summary","headline","about me"],
        "education":     ["education","degree","college","university","cgpa","gpa","marks"],
        "certifications":["certification","certificate","course","udemy","coursera","badge"],
        "length":        ["length","too short","too long","word count","pages","one page","how long"],
        "format":        ["format","template","design","layout","font","margin","ats format"],
        "keywords":      ["keyword","buzzword","jd keyword"],
        "networking":    ["network","referral","connect","reach out","cold email","hiring manager"],
        "improvement":   ["how to improve","make better","what to fix","fix resume","tips","suggestion"],
        "linkedin_post": ["linkedin post","post for linkedin","generate post","write post"],
        "roast":         ["roast","be honest","brutal","harsh","trash","bad","terrible","honest feedback"],
        "interview_q":   ["interview question","what questions","what will they ask","question they ask","practice question"],
        "before_after":  ["before after","compare","improvement","how much better","score improved"],
    }
    for intent, triggers in intents.items():
        if any(t in m for t in triggers):
            return intent
    return "unknown"


def smart_chatbot_reply(user_msg: str, resume: str, skill_data: dict) -> str:
    intent = detect_intent(user_msg)
    found   = [s for d in skill_data.values() for s in d["found"]]
    missing = [s for d in skill_data.values() for s in d["missing"]]
    tl = resume.lower()
    wc = len(resume.split())

    if intent == "greeting":
        return f"Hey! 👋 I'm ResumeIQ — your personal resume coach. Scanned your resume and found **{len(found)} skills**: {', '.join(found[:6]) or 'none yet'}. Ask me anything!"

    if intent == "strengths":
        if found: return f"Your strongest areas: **{', '.join(found[:7])}**.\n\nMake sure all appear in a clearly labelled **Skills** section — ATS scans for exact keywords."
        return "I couldn't detect many skills. Add a dedicated **Skills** section listing all your tools explicitly."

    if intent == "weaknesses":
        if missing: return f"Key gaps: **{', '.join(missing[:5])}**.\n\nStart with **{missing[0]}** — it's in demand across most tech roles. Check the 🚀 Career Pathfinder page for free course links!"
        return "Skill coverage looks decent! Focus on **measurable results (%)** and a strong Projects section."

    if intent == "role_fit":
        scores = {r: int(sum(1 for s in req if s in tl)/len(req)*100) for r,req in ROLES.items()}
        sorted_r = sorted(scores.items(), key=lambda x: -x[1])
        best,bs = sorted_r[0]; second,ss = sorted_r[1]
        extra = "\n\n🏢 For **FAANG**, you also need strong **DSA** (LeetCode medium/hard) + system design. Skills alone aren't enough." if any(x in user_msg.lower() for x in ["google","faang","amazon","microsoft"]) else ""
        return f"Best role fits based on your skills:\n\n🥇 **{best}** — {bs}% match\n🥈 **{second}** — {ss}% match\n\nVisit the 🚀 Career Pathfinder for full breakdown!{extra}"

    if intent == "ats":
        ats_sc, issues = run_ats(resume)
        if issues: return f"Your ATS score is **{ats_sc}/100**.\n\nTop issue: _{issues[0][0]}_\n\nFix these on the 🏠 Resume Intelligence page."
        return f"ATS score: **{ats_sc}/100** — great! Should pass most automated filters."

    if intent == "projects":
        if "project" not in tl: return "⚠️ No Projects section detected!\n\nEach project should include:\n- **What** you built\n- **Tech stack** (listed explicitly)\n- **Outcome** (e.g. '92% accuracy', '30% faster')\n\nFormat: *'Built X using Y, achieving Z.'*"
        return "✅ You have a Projects section! Make sure each entry has: what you built, tech stack, and a measurable outcome. Link GitHub repos if public."

    if intent == "interview":
        sq = found[0] if found else "your primary language"
        return f"Interview prep roadmap:\n\n**DSA**: LeetCode easy/medium — arrays, strings, hashmaps\n**Technical**: Code confidently in {sq}, know your projects inside-out\n**HR**: 'Tell me about yourself', 'Why this company?', 'Strengths/weaknesses'\n\nFor every project: *what, why, how, result* — practice this answer."

    if intent == "salary":
        if any(s in found for s in ["tensorflow","pytorch","scikit-learn","nlp"]): return "With ML/AI skills: **₹5–12 LPA** at product companies. FAANG-tier: **₹20–40 LPA** for exceptional profiles."
        if any(s in found for s in ["aws","docker","kubernetes"]): return "With Cloud/DevOps skills: **₹5–12 LPA** fresher range. AWS-certified candidates get better packages."
        if any(s in found for s in ["react","javascript","node.js"]): return "With Frontend/Full-Stack: **₹3–8 LPA** at startups. Senior React roles go much higher."
        return "Fresher packages: **₹3–6 LPA** at service companies, **₹8–20 LPA** at product companies. Skills + DSA determine the bracket."

    if intent == "cover_letter":
        return "Head to the ✉️ Cover Letter page! Fill in name, role, and company — I'll generate a personalised letter using your actual resume skills."

    if intent == "linkedin":
        parts = []
        if "linkedin" not in tl: parts.append("**LinkedIn URL missing** — add your profile link.")
        if "github"   not in tl: parts.append("**GitHub URL missing** — very important for tech roles. Pin your best repos.")
        return ("✅ Great — you have both LinkedIn and GitHub!" if not parts else "🔗 " + "\n\n".join(parts))

    if intent == "fresher":
        return "As a fresher, what matters most:\n\n1. **Projects** — 2–3 strong projects beat a blank Experience section\n2. **Skills section** — list all tools explicitly\n3. **GitHub** — pin your best repos\n4. **Certifications** — Coursera/Udemy certs show initiative\n5. **Internships** — even 1 month counts\n\nFocus on product companies — they value skills over years."

    if intent == "skills_list":
        if found: return f"✅ **Found ({len(found)}):** {', '.join(found)}\n\n❌ **Not detected ({len(missing[:6])}):** {', '.join(missing[:6])}\n\nMake sure skills are in a dedicated **Skills** section."
        return "No skills clearly detected. Add a **Skills** section with all your tools and technologies listed explicitly."

    if intent == "action_verbs":
        return "Strong action verbs for bullet points:\n\n✅ **Built · Developed · Designed · Led · Implemented · Deployed · Automated · Reduced · Increased · Optimised**\n\n❌ Avoid: 'Was responsible for', 'Helped with', 'Worked on'\n\n**Format**: *Verb + What + Result* → *'Developed REST API using Django, reducing response time by 40%'*"

    if intent == "summary":
        s3 = ', '.join(found[:3]) if found else 'Python and web development'
        return f"A good summary is **2–3 lines max** at the top:\n\n*'CS graduate with hands-on experience in {s3}. Built real-world projects solving measurable problems. Seeking a [target role] at a product-driven company.'*\n\nKeep it specific — generic summaries get ignored."

    if intent == "education":
        return "Education format:\n\n- Degree, Major — College Name (Year)\n- CGPA: X.X / 10 *(include only if above 7.0)*\n- Relevant coursework *(optional)*\n- Academic achievements *(add these!)*"

    if intent == "certifications":
        certs = [f"- [{n}]({u}) — {p}" for s,(n,p,u) in COURSES.items() if s in missing[:5]]
        return "Based on your gaps, I'd recommend:\n" + ("\n".join(certs) if certs else "You're looking good!") + "\n\nFree options: Google's Python Course (Coursera), CS50 (Harvard), freeCodeCamp."

    if intent == "length":
        if wc < 300:   return f"Resume is ~{wc} words — **too short**. Aim for **400–600 words**. Expand project descriptions."
        elif wc > 900: return f"Resume is ~{wc} words — **possibly too long**. Freshers should keep to **1 page**. Trim vague bullets."
        return f"Resume is ~{wc} words — **perfect length!** One page is ideal."

    if intent == "format":
        return "ATS-safe format:\n\n✅ Standard headers: Skills, Experience, Education, Projects\n✅ Single-column layout\n✅ Font: Calibri/Arial 10–12pt\n✅ Save as PDF\n\n❌ Avoid: tables, text boxes, graphics\n\nUse the 📝 Resume Builder page for a pre-formatted template!"

    if intent == "keywords":
        return "Keyword strategy:\n\n1. Copy keywords directly from the JD\n2. Add naturally in Skills and Experience\n3. Match exact spelling as in the JD\n4. Use 🎯 JD Matcher to see what you're missing!"

    if intent == "networking":
        return "Networking tips:\n\n1. **LinkedIn alumni search** — find people from your college at target companies\n2. **Cold message**: *'Hi [Name], I'm a fresher interested in [company]. Would love 10 mins to hear about your experience there.'*\n3. **Referrals** — moves you to the top of the pile\n4. **GitHub** — open source contributions get noticed"

    if intent == "linkedin_post":
        return f"Here's a LinkedIn post template:\n\n---\n🚀 Excited to share my latest project!\n\nBuilt [PROJECT NAME] using {', '.join(found[:4]) if found else 'Python and Streamlit'}.\n\n🔧 **What it does**: [1–2 line description]\n📊 **Key result**: [metric — 'improved X by Y%']\n\nThis taught me a lot about {found[0] if found else 'software development'}.\n\nCheck it out → [GitHub link]\n\n#OpenToWork #Python #BuildInPublic #FresherDeveloper\n---\n\n*(Fill in the blanks — posts with metrics get 3× more engagement!)*"

    if intent == "improvement":
        tips = smart_feedback(resume, skill_data, 0)
        return "Top improvements:\n\n" + "\n\n".join(tips)

    if intent == "roast":
        return "😤 Alright, you asked for it. Head to the **🔥 Resume Roast** section on the Dashboard page and click the Roast button. I won't hold back!"

    if intent == "interview_q":
        qs = get_interview_questions(skill_data)
        lines = [f"**{i+1}.** [{q[0].upper()}] {q[1]}" for i,q in enumerate(qs[:5])]
        return "Interview questions based on your resume:\n\n" + "\n\n".join(lines) + "\n\nSee all questions on the 📝 Interview Prep page!"

    if intent == "before_after":
        hist = st.session_state.get("score_history",[])
        if len(hist) >= 2:
            old, new = hist[0][1], hist[-1][1]
            diff = new - old
            return f"Your resume score went from **{old}** → **{new}** ({'+' if diff>=0 else ''}{diff} points). {'Great improvement! 🎉' if diff > 0 else 'Try applying the AI feedback tips to improve!'}"
        return "Upload 2 different versions of your resume to see the before/after comparison. Head to the 📊 Score History section on the Dashboard!"

    # Fallback
    scores = {r: int(sum(1 for s in req if s in tl)/len(req)*100) for r,req in ROLES.items()}
    best   = max(scores, key=scores.get)
    return (f"Quick snapshot of your resume:\n\n"
            f"🔧 **Skills ({len(found)}):** {', '.join(found[:5]) or 'None clearly listed'}\n"
            f"🎯 **Best role:** {best} ({scores[best]}%)\n"
            f"⚠️ **Top gaps:** {', '.join(missing[:4]) or 'None major'}\n"
            f"📝 **Word count:** ~{wc} words\n\n"
            f"Try: *'What are my strengths?'*, *'Which job suits me?'*, *'Roast my resume'*, *'Give interview questions'*, or *'What salary can I expect?'*")


def generate_cover_letter(name, role, company, tone, resume, skill_data):
    found = [s for d in skill_data.values() for s in d["found"]]
    skill_str = ", ".join(found[:5]) if found else "various technical skills"
    tl = resume.lower()
    bg = ("professional experience and projects" if "experience" in tl and "project" in tl
          else "hands-on projects" if "project" in tl else "academic coursework")
    openings = {
        "Professional": f"I am writing to express my strong interest in the {role} position at {company}.",
        "Enthusiastic": f"I am genuinely excited to apply for the {role} role at {company} — a company I deeply respect!",
        "Concise":      f"I am applying for the {role} position at {company} and believe my profile is a strong fit.",
    }
    return f"""Dear Hiring Manager,

{openings.get(tone, openings['Professional'])} With a solid foundation in {skill_str}, I am confident in my ability to contribute meaningfully from day one.

Through my {bg}, I have developed strong problem-solving abilities and hands-on experience with technologies relevant to this role. I have applied skills like {skill_str} in real projects, consistently delivering results in collaborative environments.

I am drawn to {company} because of its focus on innovation and impact. I would be thrilled to bring my skills and drive to your team.

Thank you for your time and consideration. I look forward to discussing how I can contribute to {company}'s goals.

Sincerely,
{name}"""


# ═══════════════════════════════════════════════════════
#  RESUME BUILDER
# ═══════════════════════════════════════════════════════

def build_resume_docx(data: dict, template: str) -> bytes:
    doc = DocxDocument()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(0.75)
        sec.left_margin = sec.right_margin = Inches(0.85)

    accents = {"Classic": RGBColor(79,70,229), "Modern": RGBColor(16,185,129), "Minimal": RGBColor(100,116,139)}
    accent  = accents.get(template, accents["Classic"])

    contact_line = "  |  ".join(x for x in [data.get("email",""), data.get("phone",""), data.get("linkedin",""), data.get("github","")] if x)

    def add_heading(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(title.upper())
        r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = accent
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(3)
        r2 = p2.add_run("─" * 72); r2.font.size = Pt(7); r2.font.color.rgb = RGBColor(200,200,200)

    def add_body(content):
        for line in content.strip().split("\n"):
            line = line.strip()
            if not line: continue
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
            r = p.add_run(line); r.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(data.get("name","Your Name"))
    r.bold = True; r.font.size = Pt(20); r.font.color.rgb = accent

    if contact_line:
        p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(6)
        r2 = p2.add_run(contact_line); r2.font.size = Pt(9); r2.font.color.rgb = RGBColor(80,80,80)

    for title, key in [("Professional Summary","summary"),("Skills","skills"),("Work Experience","experience"),("Education","education"),("Projects","projects"),("Certifications","certifications")]:
        content = data.get(key,"")
        if content and content.strip():
            add_heading(title); add_body(content)

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════
#  SESSION STATE INIT
# ═══════════════════════════════════════════════════════
for key, val in [("started",False),("resume_text",None),("filename",None),
                  ("chat",[]),("score_history",[]),("template_choice","Classic")]:
    if key not in st.session_state:
        st.session_state[key] = val


# ═══════════════════════════════════════════════════════
#  LANDING PAGE
# ═══════════════════════════════════════════════════════
if not st.session_state.started:
    st.markdown("""
    <div class="landing-wrap">
      <div class="landing-chip">AI Resume Analyzer</div>
      <h1 class="landing-title">📄 ResumeIQ</h1>
      <p class="landing-sub">Upload your resume and get instant analysis — ATS score, skill gaps, career matching and more.</p>
    </div>
    """, unsafe_allow_html=True)

    c1,c2,c3 = st.columns([1,2,1])
    with c2:
        if st.button("🚀  Get Started ", use_container_width=True):
            st.session_state.started = True
            st.rerun()

    st.markdown("""
    <div class="landing-features">
      <div class="landing-feat"><div class="landing-feat-icon">🏠</div><div class="landing-feat-label">Resume Intelligence</div></div>
      <div class="landing-feat"><div class="landing-feat-icon">🎯</div><div class="landing-feat-label">JD Matcher</div></div>
      <div class="landing-feat"><div class="landing-feat-icon">🚀</div><div class="landing-feat-label">Career Pathfinder</div></div>
      <div class="landing-feat"><div class="landing-feat-icon">💬</div><div class="landing-feat-label">AI Assistant</div></div>
      <div class="landing-feat"><div class="landing-feat-icon">✉️</div><div class="landing-feat-label">Cover Letter</div></div>
      <div class="landing-feat"><div class="landing-feat-icon">📝</div><div class="landing-feat-label">Resume Builder</div></div>
      <div class="landing-feat"><div class="landing-feat-icon">🔥</div><div class="landing-feat-label">Resume Roast</div></div>
      <div class="landing-feat"><div class="landing-feat-icon">🎤</div><div class="landing-feat-label">Interview Prep</div></div>
    </div>
    """, unsafe_allow_html=True)
    st.stop()


# ═══════════════════════════════════════════════════════
#  SIDEBAR
# ═══════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("## 📄 ResumeIQ")
    st.markdown("---")
    uploaded = st.file_uploader("Upload Resume (PDF / DOCX)", type=["pdf","docx"])
    if uploaded:
        if st.session_state.filename != uploaded.name:
            text = extract_text(uploaded)
            st.session_state.resume_text = text
            st.session_state.filename    = uploaded.name
            st.session_state.chat        = []
            if not text.startswith("[ERROR"):
                score = compute_score(text)
                st.session_state.score_history.append((datetime.now().strftime("%H:%M"), score, uploaded.name))
        st.success(f"✅ {uploaded.name}")

    st.markdown("---")
    page = st.radio("Navigate", [
        "🏠 Resume Intelligence",
        "🎯 JD Matcher",
        "🚀 Career Pathfinder",
        "💬 AI Assistant",
        "✉️ Cover Letter",
        "📝 Resume Builder",
        "🎤 Interview Prep",
    ])
    st.markdown("---")
    if st.button("← Back to Home"):
        st.session_state.started = False; st.rerun()
    st.caption("ResumeIQ")


# ═══════════════════════════════════════════════════════
#  HERO
# ═══════════════════════════════════════════════════════
st.markdown("""
<div class="hero">
  <div class="hero-chip">✦ AI-POWERED RESUME TOOL</div>
  <h1>📄 ResumeIQ</h1>
  <p class="hero-sub">ATS Analysis · Skill Intelligence · Career Matching · Interview Prep · Resume Builder</p>
  </div>
</div>
""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════
#  RESUME BUILDER (no resume needed)
# ═══════════════════════════════════════════════════════
if page == "📝 Resume Builder":
    st.markdown('<p class="sec-title">📝 Resume Builder — Build ATS-Friendly Resume</p>', unsafe_allow_html=True)

    # Template selector
    st.markdown('<div class="card"><div class="sec-title">Step 1 — Choose Template</div>', unsafe_allow_html=True)
    tc = st.session_state.template_choice
    t1,t2,t3 = st.columns(3)
    for col,(tname,icon,desc) in zip([t1,t2,t3],[
        ("Classic","📄","Traditional layout. Best for MNCs & corporates."),
        ("Modern","✨","Bold accents. Best for startups & tech roles."),
        ("Minimal","🔲","Ultra clean. Best for product companies."),
    ]):
        with col:
            sel = " tmpl-selected" if tc==tname else ""
            st.markdown(f'<div class="tmpl-card{sel}"><div style="font-size:2rem">{icon}</div><strong style="color:#c4b5fd">{tname}</strong><p style="font-size:.78rem;color:#64748b;margin:6px 0">{desc}</p></div>', unsafe_allow_html=True)
            if st.button(f"Select {tname}", use_container_width=True, key=f"t_{tname}"):
                st.session_state.template_choice = tname; st.rerun()
    st.markdown(f'<p style="color:#34d399;font-size:.83rem;margin-top:6px">✅ Selected: <strong>{tc}</strong></p></div>', unsafe_allow_html=True)

    st.markdown('<div class="card"><div class="sec-title">Step 2 — Fill Details</div>', unsafe_allow_html=True)
    st.caption("👇 Click below to auto-fill a complete sample fresher profile for your demo — one click!")
    if st.button("⚡ Load Sample Data (Demo Mode)"):
        for k,v in SAMPLE_DATA.items(): st.session_state[f"rb_{k}"] = v
        st.rerun()
    st.markdown("</div>", unsafe_allow_html=True)

    c1,c2 = st.columns(2)
    with c1:
        name  = st.text_input("Full Name *",  key="rb_name",  placeholder="Rahul Sharma")
        email = st.text_input("Email *",       key="rb_email", placeholder="rahul@gmail.com")
        phone = st.text_input("Phone",         key="rb_phone", placeholder="+91 98765 43210")
    with c2:
        lin   = st.text_input("LinkedIn",      key="rb_linkedin", placeholder="linkedin.com/in/rahulsharma")
        gh    = st.text_input("GitHub",        key="rb_github",   placeholder="github.com/rahulsharma")
    summary  = st.text_area("Professional Summary",     key="rb_summary",       height=80,  placeholder="Final year CS student skilled in Python and ML...")
    skills   = st.text_area("Skills (comma separated)", key="rb_skills",        height=65,  placeholder="Python, SQL, TensorFlow, React, Docker, Git")
    exp      = st.text_area("Work Experience",           key="rb_experience",    height=120, placeholder="ML Intern — XYZ Tech (Jun–Aug 2024)\n• Built REST APIs, reducing response time by 30%")
    edu      = st.text_area("Education",                 key="rb_education",     height=75,  placeholder="B.Tech CSE — VIT University (2021–2025) | CGPA: 8.4/10")
    projects = st.text_area("Projects",                  key="rb_projects",      height=120, placeholder="Resume Analyzer (Python, Streamlit)\n• Built ATS checker used by 200+ students")
    certs    = st.text_area("Certifications",            key="rb_certifications",height=65,  placeholder="Deep Learning Specialization — Coursera (2024)")

    st.markdown("---")
    if st.button("🚀 Build & Download Resume", use_container_width=True):
        rbn  = st.session_state.get("rb_name","")
        rbe  = st.session_state.get("rb_email","")
        if not rbn or not rbe:
            st.error("Please fill in at least Name and Email.")
        else:
            form_data = {"name":name,"email":email,"phone":phone,"linkedin":lin,"github":gh,
                         "summary":summary,"skills":skills,"experience":exp,"education":edu,
                         "projects":projects,"certifications":certs}
            resume_bytes = build_resume_docx(form_data, tc)
            st.success("✅ Resume built successfully!")
            st.download_button("📥 Download Resume (.docx)", data=resume_bytes,
                file_name=f"{rbn.replace(' ','_')}_Resume_{tc}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True)
            # Quick ATS check
            preview = " ".join(filter(None,[summary,skills,exp,edu,projects,certs,rbn,rbe,phone,lin]))
            preview += " experience education skills project"
            built_ats,built_issues = run_ats(preview)
            st.markdown(f'<div class="card"><div class="sec-title">✅ ATS Check — Built Resume ({built_ats}/100)</div>', unsafe_allow_html=True)
            if not built_issues: st.success("Passes all ATS checks!")
            else:
                for msg,_ in built_issues: st.warning(f"⚠️ {msg}")
            st.markdown("</div>", unsafe_allow_html=True)
    st.stop()


# ═══════════════════════════════════════════════════════
#  GUARD
# ═══════════════════════════════════════════════════════
resume = st.session_state.resume_text
if not resume:
    st.markdown("""
    <div class="card" style="text-align:center;padding:50px">
      <div style="font-size:3rem">📎</div>
      <p style="color:#94a3b8;font-size:1rem;margin-top:12px">Upload your resume in the sidebar to get started.<br>
      <span style="font-size:.85rem;color:#475569">Or go to 📝 Resume Builder to create one from scratch.</span></p>
    </div>""", unsafe_allow_html=True)
    st.stop()
if resume.startswith("[ERROR"):
    st.error(resume); st.stop()

skill_data  = get_skills(resume)
ats_sc, issues = run_ats(resume)
all_found  = sum(len(d["found"]) for d in skill_data.values())
all_total  = sum(len(d["found"])+len(d["missing"]) for d in skill_data.values())
skill_pct  = int(all_found/max(all_total,1)*100)
final_sc   = int(ats_sc*0.5 + skill_pct*0.5)


# ═══════════════════════════════════════════════════════
#  PAGE: RESUME INTELLIGENCE (Dashboard)
# ═══════════════════════════════════════════════════════
if page == "🏠 Resume Intelligence":
    st.markdown('<p class="sec-title">🏠 Resume Intelligence</p>', unsafe_allow_html=True)

    # Gauge charts
    def gauge(val, title, color):
        fig = go.Figure(go.Indicator(
            mode="gauge+number", value=val,
            title={"text":title,"font":{"size":13,"family":"Inter"}},
            number={"suffix":"%","font":{"size":26}},
            gauge={"axis":{"range":[0,100],"tickwidth":1,"tickcolor":"#2a2a3e"},
                   "bar":{"color":color,"thickness":0.28},"bgcolor":"#11111e","borderwidth":0,
                   "steps":[{"range":[0,40],"color":"#1f0a0a"},{"range":[40,70],"color":"#1c1200"},{"range":[70,100],"color":"#052e16"}],
                   "threshold":{"line":{"color":color,"width":3},"thickness":.8,"value":val}}))
        fig.update_layout(height=200, margin=dict(t=30,b=5,l=15,r=15), paper_bgcolor="#11111e", font={"family":"Inter","color":"#e2e8f0"})
        return fig

    g1,g2,g3 = st.columns(3)
    for col,val,title,color in [(g1,final_sc,"⭐ Overall",  "#7c3aed"),(g2,ats_sc,"🤖 ATS Score","#059669"),(g3,skill_pct,"🔧 Skills","#d97706")]:
        with col:
            st.markdown('<div class="card" style="padding:10px">', unsafe_allow_html=True)
            st.plotly_chart(gauge(val,title,color), use_container_width=True)
            st.markdown("</div>", unsafe_allow_html=True)

    # Benchmarking chart
    st.markdown('<div class="card"><div class="sec-title">📊 How You Compare</div>', unsafe_allow_html=True)
    bf = go.Figure()
    bf.add_trace(go.Bar(name="You",            x=["Overall","ATS","Skills"], y=[final_sc,ats_sc,skill_pct], marker_color="#a78bfa", text=[f"{v}%" for v in [final_sc,ats_sc,skill_pct]], textposition="outside", textfont=dict(color="#ffffff",size=12)))
    bf.add_trace(go.Bar(name="Avg Fresher",    x=["Overall","ATS","Skills"], y=[52,60,40],  marker_color="#f97316", text=["52%","60%","40%"],  textposition="outside", textfont=dict(color="#ffffff",size=12)))
    bf.add_trace(go.Bar(name="Good Candidate", x=["Overall","ATS","Skills"], y=[72,80,65],  marker_color="#34d399", text=["72%","80%","65%"],  textposition="outside", textfont=dict(color="#ffffff",size=12)))
    bf.update_layout(barmode="group", height=260, yaxis=dict(range=[0,130],gridcolor="#1e1e2e",tickfont=dict(color="#e2e8f0")),
                     xaxis=dict(tickfont=dict(color="#e2e8f0",size=13)),
                     plot_bgcolor="#11111e", paper_bgcolor="#11111e", font=dict(color="#e2e8f0"),
                     margin=dict(t=20,b=10), legend=dict(bgcolor="#1e1e2e",bordercolor="#2a2a3e",font=dict(color="#e2e8f0")))
    st.plotly_chart(bf, use_container_width=True)
    st.markdown("</div>", unsafe_allow_html=True)

    # Score history
    hist = st.session_state.score_history
    if len(hist) >= 2:
        st.markdown('<div class="card"><div class="sec-title">📈 Score History — Before vs After</div>', unsafe_allow_html=True)
        hf = go.Figure(go.Scatter(
            x=[h[0] + " • " + h[2][:15] for h in hist], y=[h[1] for h in hist],
            mode="lines+markers+text", text=[str(h[1]) for h in hist],
            textposition="top center", line=dict(color="#7c3aed",width=3),
            marker=dict(size=10,color="#34d399"),
        ))
        hf.update_layout(height=220, plot_bgcolor="#11111e", paper_bgcolor="#11111e",
                         font=dict(color="#e2e8f0"), margin=dict(t=10,b=10),
                         yaxis=dict(range=[0,105],gridcolor="#1e1e2e"))
        st.plotly_chart(hf, use_container_width=True)
        st.markdown("</div>", unsafe_allow_html=True)

    # ATS Checklist
    st.markdown('<div class="card"><div class="sec-title">📋 ATS Checklist</div>', unsafe_allow_html=True)
    if not issues: st.success("✅ Your resume passes all ATS checks!")
    else:
        for msg,_ in issues: st.markdown(f'<div style="color:#fca5a5;font-size:.88rem;padding:4px 0">❌ {msg}</div>', unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)

    # Skill grid
    st.markdown('<div class="card"><div class="sec-title">🔍 Skills by Category</div>', unsafe_allow_html=True)
    sc1,sc2 = st.columns(2)
    for i,(cat,data) in enumerate(skill_data.items()):
        col = sc1 if i%2==0 else sc2
        tot = len(data["found"])+len(data["missing"])
        pct = int(len(data["found"])/max(tot,1)*100)
        with col:
            fp = "".join(f'<span class="pill have">✔ {s}</span>' for s in data["found"])
            mp = "".join(f'<span class="pill missing">✗ {s}</span>' for s in data["missing"])
            st.markdown(f"""
            <div class="skill-cat-card">
              <div class="skill-cat-title">{cat}</div>
              <div class="skill-mini-prog-wrap"><div class="skill-mini-prog" style="width:{pct}%;background:{prog_color(pct)}"></div></div>
              <span style="font-size:.7rem;color:#64748b">{len(data['found'])}/{tot} — {pct}%</span>
              <div style="margin-top:7px">{fp}{mp}</div>
            </div>""", unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)

    # AI Feedback
    st.markdown('<div class="card"><div class="sec-title">🤖 Smart AI Feedback</div>', unsafe_allow_html=True)
    if st.button("Get AI Feedback"):
        for tip in smart_feedback(resume, skill_data, ats_sc):
            st.markdown(f'<div class="tip-box">{tip}</div>', unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)

    # 🔥 ROAST MODE
    st.markdown('<div class="card"><div class="sec-title">🔥 Resume Roast Mode</div>', unsafe_allow_html=True)
    st.caption("Brutally honest feedback — only if you can handle the truth 😅")
    if st.button("🔥 Roast My Resume"):
        roasts = generate_roast(resume, skill_data, ats_sc)
        for r in roasts:
            st.markdown(f'<div class="roast-box">{r}</div>', unsafe_allow_html=True)
        st.markdown('<div class="tip-box-green" style="margin-top:12px">💡 Roast aside — these are real issues. Fix them and your resume will be genuinely stronger.</div>', unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)

    # Stats
    wc = len(resume.split())
    st.markdown(f"""
    <div class="card">
      <div class="sec-title">📄 Document Stats</div>
      <div style="display:grid;grid-template-columns:repeat(3,1fr);gap:12px;text-align:center">
        <div><div style="font-size:1.4rem;font-weight:700;color:#a78bfa">{wc}</div><div style="font-size:.72rem;color:#475569">Words</div></div>
        <div><div style="font-size:1.4rem;font-weight:700;color:#34d399">{all_found}</div><div style="font-size:.72rem;color:#475569">Skills Found</div></div>
        <div><div style="font-size:1.4rem;font-weight:700;color:#fbbf24">{len(issues)}</div><div style="font-size:.72rem;color:#475569">ATS Issues</div></div>
      </div>
    </div>""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════
#  PAGE: JD MATCHER
# ═══════════════════════════════════════════════════════
elif page == "🎯 JD Matcher":
    st.markdown('<p class="sec-title">🎯 JD Matcher — Resume vs Job Description</p>', unsafe_allow_html=True)
    jd = st.text_area("Paste the Job Description here", height=180, placeholder="Copy-paste the full JD here…")

    if jd:
        sw  = {"the","and","for","are","with","have","that","this","will","you","from","our","their","we","a","an","in","of","to","is","be","as","on","at","by","or","it","its","your","they","do","not","has","was","were","but","if","all","can","which","who","what"}
        rw  = set(re.findall(r'\b[a-zA-Z][a-zA-Z0-9+#.]*\b', resume.lower())) - sw
        jraw= re.findall(r'\b[a-zA-Z][a-zA-Z0-9+#.]*\b', jd.lower())
        jfreq = Counter(w for w in jraw if w not in sw and len(w)>2)
        jw  = set(jfreq.keys())
        matched = sorted(rw & jw)
        missing = sorted(jw - rw, key=lambda w:-jfreq[w])
        score   = int(len(matched)/max(len(jw),1)*100)

        c1,c2,c3 = st.columns(3)
        fc_score = "#34d399" if score>=70 else ("#fbbf24" if score>=40 else "#f87171")
        c1.markdown(f'''<div style="background:#11111e;border:1px solid #1e1e30;border-radius:12px;padding:16px;text-align:center">
            <div style="font-size:1.8rem;font-weight:800;color:{fc_score}">{score}%</div>
            <div style="font-size:.75rem;color:#64748b;margin-top:4px">Match Score</div></div>''', unsafe_allow_html=True)
        c2.markdown(f'''<div style="background:#11111e;border:1px solid #1e1e30;border-radius:12px;padding:16px;text-align:center">
            <div style="font-size:1.8rem;font-weight:800;color:#1a1f2e">{len(matched)}</div>
            <div style="font-size:.75rem;color:#64748b;margin-top:4px">Keywords Matched</div></div>''', unsafe_allow_html=True)
        c3.markdown(f'''<div style="background:#11111e;border:1px solid #1e1e30;border-radius:12px;padding:16px;text-align:center">
            <div style="font-size:1.8rem;font-weight:800;color:#1a1f2e">{len(missing)}</div>
            <div style="font-size:.75rem;color:#64748b;margin-top:4px">Keywords Missing</div></div>''', unsafe_allow_html=True)
        st.markdown(f'<div class="prog-wrap"><div class="prog-fill" style="width:{score}%;background:{prog_color(score)}"></div></div>', unsafe_allow_html=True)

        col1,col2 = st.columns(2)
        with col1:
            st.markdown('<div class="card">', unsafe_allow_html=True)
            st.markdown(f'<div class="sec-title">✅ Matched ({len(matched)})</div>', unsafe_allow_html=True)
            if matched:
                colors = ["#d1fae5","#a7f3d0","#6ee7b7","#34d399","#10b981"]
                cloud  = "".join(f'<span class="wcloud-word" style="font-size:{min(0.85+jfreq.get(w,1)*0.1,1.3)}rem;background:{colors[i%5]}22;color:{colors[i%5]};border:1px solid {colors[i%5]}55">{w}</span>' for i,w in enumerate(matched[:40]))
                st.markdown(cloud, unsafe_allow_html=True)
            else: st.info("No matched keywords.")
            st.markdown("</div>", unsafe_allow_html=True)

        with col2:
            st.markdown('<div class="card">', unsafe_allow_html=True)
            st.markdown(f'<div class="sec-title">❌ Missing ({len(missing)}) — Priority Ranked</div>', unsafe_allow_html=True)
            if missing:
                max_f = jfreq[missing[0]] if missing else 1
                rows  = [missing[i:i+4] for i in range(0,min(len(missing),24),4)]
                for row in rows:
                    rh = ""
                    for w in row:
                        ratio = jfreq[w]/max(max_f,1)
                        b = '<span class="badge badge-high">HIGH</span>' if ratio>=.6 else ('<span class="badge badge-medium">MED</span>' if ratio>=.3 else '<span class="badge badge-low">LOW</span>')
                        rh += f'{b}<span class="pill missing">{w}</span>&nbsp;'
                    st.markdown(f'<div style="margin:4px 0">{rh}</div>', unsafe_allow_html=True)
            else: st.success("No missing keywords — great match!")
            st.markdown("</div>", unsafe_allow_html=True)

        st.markdown('<div class="card"><div class="sec-title">🤖 Smart Analysis</div>', unsafe_allow_html=True)
        if st.button("Analyse Match"):
            for tip in smart_jd_analysis(matched, missing[:10], score, jd):
                st.markdown(f'<div class="tip-box">{tip}</div>', unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════
#  PAGE: CAREER PATHFINDER
# ═══════════════════════════════════════════════════════
elif page == "🚀 Career Pathfinder":
    st.markdown('<p class="sec-title">🚀 Career Pathfinder</p>', unsafe_allow_html=True)
    tl = resume.lower()
    role_scores = {r: int(sum(1 for s in req if s in tl)/len(req)*100) for r,req in ROLES.items()}

    fig = go.Figure(go.Pie(labels=list(role_scores.keys()), values=list(role_scores.values()), hole=0.42,
                           marker_colors=["#7c3aed","#059669","#d97706","#ef4444","#0ea5e9"],
                           textfont=dict(color="#ffffff",size=13),
                                         textinfo="label+percent"))
    fig.update_layout(height=300, paper_bgcolor="#1e1e30", font=dict(color="#e2e8f0"),
                      margin=dict(t=10,b=10), legend=dict(bgcolor="#1e1e30",font=dict(color="#e2e8f0")))
    st.plotly_chart(fig, use_container_width=True)

    for role,score in sorted(role_scores.items(), key=lambda x:-x[1]):
        req   = ROLES[role]
        found = [s for s in req if s in tl]
        miss  = [s for s in req if s not in tl]
        fc    = "#34d399" if score>=70 else ("#fbbf24" if score>=40 else "#f87171")
        st.markdown(f"""
        <div class="card">
          <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:8px">
            <span style="font-weight:700;font-size:1rem;color:#e2e8f0">{role}</span>
            <span style="color:{fc};font-weight:800;font-size:1rem;background:{fc}22;padding:4px 12px;border-radius:20px;border:1px solid {fc}55">{score}%</span>
          </div>
          <div class="prog-wrap"><div class="prog-fill" style="width:{score}%;background:{prog_color(score)}"></div></div>
          <div style="margin-top:10px">{"".join(f'<span class="pill have">{s}</span>' for s in found)}{"".join(f'<span class="pill missing">{s}</span>' for s in miss)}</div>
        </div>""", unsafe_allow_html=True)
        for s in miss:
            if s in COURSES:
                n,p,u = COURSES[s]
                st.caption(f"📚 [{n}]({u}) — {p}")


# ═══════════════════════════════════════════════════════
#  PAGE: AI ASSISTANT (Chatbot)
# ═══════════════════════════════════════════════════════
elif page == "💬 AI Assistant":
    st.markdown('<p class="sec-title">💬 AI Assistant</p>', unsafe_allow_html=True)
    st.caption("Knows your resume inside out. Ask anything in your own words — even casually.")

    if "chat" not in st.session_state: st.session_state.chat = []

    # Quick buttons
    st.markdown('<div class="sec-title" style="margin-top:8px">Quick Questions</div>', unsafe_allow_html=True)
    quick_btns = [
        ("💪 Strengths",       "What are my strengths?"),
        ("⚠️ Weaknesses",      "What are my weaknesses?"),
        ("🎯 Best Role",       "Which job role suits me best?"),
        ("📈 How to Improve",  "How can I improve my resume?"),
        ("🎤 Interview Tips",  "Give me interview tips."),
        ("💰 Salary",          "What salary can I expect?"),
        ("📝 Resume Length",   "Is my resume the right length?"),
        ("🏆 Certifications",  "What certifications should I do?"),
        ("🎓 Fresher Tips",    "I am a fresher, what should I focus on?"),
        ("🌐 LinkedIn Post",   "Generate a LinkedIn post for me."),
        ("🔠 Action Verbs",    "What action verbs should I use?"),
        ("🤖 ATS Score",       "What is my ATS score?"),
    ]
    quick = None
    cols  = st.columns(4)
    for i,(label,q) in enumerate(quick_btns):
        if cols[i%4].button(label, key=f"qb{i}", use_container_width=True): quick = q

    # Chat window
    st.markdown('<div class="chat-wrap">', unsafe_allow_html=True)
    if not st.session_state.chat:
        st.markdown('<p style="text-align:center;color:#334155;font-size:.85rem;padding:20px 0">👋 Pick a quick question above or type anything below!</p>', unsafe_allow_html=True)
    for sender,msg in st.session_state.chat:
        if sender == "user":
            st.markdown(f'<div class="msg-row-user"><div class="bubble-user">{msg}</div><div class="avatar avatar-user">🧑</div></div>', unsafe_allow_html=True)
        else:
            st.markdown(f'<div class="msg-row-bot"><div class="avatar avatar-bot">🤖</div><div class="bubble-bot">{msg.replace(chr(10),"<br>")}</div></div>', unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)

    user_input  = st.chat_input("Type your question… (e.g. Am I ready for Google?)")
    final_input = quick or user_input

    if final_input:
        with st.empty():
            st.markdown('<div class="msg-row-bot"><div class="avatar avatar-bot">🤖</div><div class="bubble-bot"><span class="typing-dot"></span><span class="typing-dot"></span><span class="typing-dot"></span></div></div>', unsafe_allow_html=True)
            time.sleep(0.7)
        reply = smart_chatbot_reply(final_input, resume, skill_data)
        st.session_state.chat.append(("user", final_input))
        st.session_state.chat.append(("bot",  reply))
        st.rerun()

    if st.session_state.chat and st.button("🗑 Clear chat"):
        st.session_state.chat = []; st.rerun()


# ═══════════════════════════════════════════════════════
#  PAGE: COVER LETTER
# ═══════════════════════════════════════════════════════
elif page == "✉️ Cover Letter":
    st.markdown('<p class="sec-title">✉️ Cover Letter Generator</p>', unsafe_allow_html=True)
    c1,c2 = st.columns(2)
    with c1:
        name    = st.text_input("Your Name",    placeholder="Rahul Sharma")
        role    = st.text_input("Job Role",     placeholder="Software Engineer")
    with c2:
        company = st.text_input("Company Name", placeholder="Google")
        tone    = st.selectbox("Tone", ["Professional","Enthusiastic","Concise"])

    if st.button("✨ Generate Cover Letter") and name and role and company:
        letter = generate_cover_letter(name, role, company, tone, resume, skill_data)
        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.text_area("Your Cover Letter", letter, height=340)
        st.download_button(
            label="⬇️ Download Cover Letter (.txt)",
            data=letter,
            file_name="cover_letter.txt",
            mime="text/plain",
            use_container_width=True,
        )
        st.markdown("</div>", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════
#  PAGE: INTERVIEW PREP
# ═══════════════════════════════════════════════════════
elif page == "🎤 Interview Prep":
    st.markdown('<p class="sec-title">🎤 Interview Prep — Questions Based on Your Resume</p>', unsafe_allow_html=True)

    found_skills = [s for d in skill_data.values() for s in d["found"]]

    st.markdown('<div class="card"><div class="sec-title">🎯 Questions Tailored to Your Skills</div>', unsafe_allow_html=True)
    st.caption(f"Based on your detected skills: {', '.join(found_skills[:6]) if found_skills else 'none detected yet'}")

    questions = get_interview_questions(skill_data)
    for i,(skill,q) in enumerate(questions):
        st.markdown(f"""
        <div class="iq-card">
          <span class="iq-num">Q{i+1}.</span>
          <span class="pill neutral">{skill}</span>&nbsp;
          {q}
        </div>""", unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)

    # HR Questions
    st.markdown('<div class="card"><div class="sec-title">🗣 Common HR Questions — With Tips</div>', unsafe_allow_html=True)
    hr_qs = [
        ("Tell me about yourself.",
         f"30-second pitch: 'I'm a [degree] student with experience in {', '.join(found_skills[:3]) if found_skills else 'Python and web development'}. I've built [X] projects and I'm now looking for a role where I can [value you bring].'"),
        ("Why should we hire you?",
         "Focus on 2–3 specific skills + a project result. 'I bring X skill, proved by Y outcome in my project. I also [soft skill example].'"),
        ("What is your greatest weakness?",
         "Pick a real but minor weakness and show you're working on it. Example: 'I tend to over-engineer solutions — I'm learning to ship faster and iterate.'"),
        ("Where do you see yourself in 5 years?",
         "Align with the company's direction. 'I want to become a strong [role] and eventually take on more responsibility in [domain].'"),
        ("Do you have any questions for us?",
         "Always say yes! Ask: 'What does a typical day look like for someone in this role?' or 'What does success look like in the first 90 days?'"),
    ]
    for q, tip in hr_qs:
        st.markdown(f'<div class="iq-card"><strong>Q: {q}</strong><br><span style="color:#94a3b8;font-size:.83rem">💡 {tip}</span></div>', unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)

    # 2-min pitch builder
    st.markdown('<div class="card"><div class="sec-title">🎙 Build Your 2-Minute Pitch</div>', unsafe_allow_html=True)
    if st.button("Generate My Pitch"):
        s3 = ', '.join(found_skills[:3]) if found_skills else 'Python and web development'
        has_proj = "project" in resume.lower()
        has_intern = "intern" in resume.lower()
        exp_line = ("I've completed an internship and built multiple real-world projects." if has_intern
                    else "I've built multiple real-world projects." if has_proj
                    else "I've developed strong technical foundations through my coursework and self-learning.")
        pitch = f"""Hi, I'm [Your Name], a final-year [Degree] student from [College].

I have hands-on experience in {s3}. {exp_line}

One project I'm particularly proud of is [your best project] — where I [what you built] using [tech stack], which resulted in [outcome or learning].

I'm now actively looking for opportunities where I can apply my technical skills and continue growing. I'm especially interested in [company/domain] because [specific reason].

I believe my background in {s3} makes me a strong fit for this role, and I'm excited about the opportunity to contribute to your team."""
        st.markdown(f'<div class="tip-box-green" style="white-space:pre-wrap;font-family:monospace;font-size:.85rem">{pitch}</div>', unsafe_allow_html=True)
        st.caption("Fill in the [brackets] with your actual details. Practice this until it's smooth and confident.")
    st.markdown("</div>", unsafe_allow_html=True)
